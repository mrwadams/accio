from typing import List, Dict, Any, Optional
import os
from pathlib import Path
import logging
from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    CSVLoader,
    UnstructuredPDFLoader,
    UnstructuredWordDocumentLoader,
    TextLoader,
    UnstructuredMarkdownLoader
)
from langchain_unstructured import UnstructuredLoader
from langchain_core.documents import Document as LangchainDocument
import pandas as pd
from docx import Document as DocxDocument
import PyPDF2
import uuid

logger = logging.getLogger(__name__)

class DocumentLoader:
    """A modular document loader that supports PDF, DOCX, and CSV files using Langchain loaders."""
    
    # Define supported file types and their configurations
    LOADER_CONFIGS = {
        '.pdf': {
            'primary_loader': PyPDFLoader,
            'fallback_loader': UnstructuredPDFLoader,
            'min_chars_per_chunk': 50,
            'max_chars_per_chunk': 8000,
            'source_range_type': 'page'
        },
        '.docx': {
            'primary_loader': Docx2txtLoader,
            'fallback_loader': UnstructuredWordDocumentLoader,
            'min_chars_per_chunk': 50,
            'max_chars_per_chunk': 8000,
            'source_range_type': 'paragraph'
        },
        '.csv': {
            'primary_loader': CSVLoader,
            'loader_kwargs': {
                'encoding': 'utf-8'
            },
            'source_range_type': 'row'
        },
        '.txt': {
            'primary_loader': TextLoader,
            'loader_kwargs': {
                'encoding': 'utf-8'
            },
            'source_range_type': 'line'
        },
        '.md': {
            'primary_loader': UnstructuredMarkdownLoader,
            'loader_kwargs': {},
            'source_range_type': 'section'
        }
    }
    
    SUPPORTED_EXTENSIONS = list(LOADER_CONFIGS.keys())
    
    @staticmethod
    def validate_content(content: str, file_type: str) -> Optional[str]:
        """
        Validate document content based on file type.
        
        Args:
            content (str): The extracted text content
            file_type (str): The file extension
            
        Returns:
            Optional[str]: Error message if validation fails, None if passes
        """
        if not content.strip():
            return "Document appears to be empty"
            
        min_chars = DocumentLoader.LOADER_CONFIGS.get(file_type, {}).get('min_chars_per_chunk', 50)
        max_chars = DocumentLoader.LOADER_CONFIGS.get(file_type, {}).get('max_chars_per_chunk', 8000)
        
        if len(content) < min_chars:
            return f"Document content too short (min {min_chars} characters)"
        
        if len(content) > max_chars * 100:  # Arbitrary large multiplier for total document
            return f"Document content exceeds maximum size"
            
        return None

    @staticmethod
    def extract_detailed_metadata(file_path: Path, schema_name: Optional[str] = None, loader_config: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Extract detailed metadata based on file type.
        
        Args:
            file_path (Path): Path to the document file
            schema_name (Optional[str]): Name of the schema this document belongs to
            loader_config (Optional[Dict[str, Any]]): Optional configuration including original_filename
            
        Returns:
            Dict[str, Any]: Detailed metadata dictionary
        """
        # Generate a unique document ID
        doc_id = str(uuid.uuid4())
        
        metadata = {
            "document_id": doc_id,
            "filename": loader_config.get("original_filename", file_path.name),  # Use original filename if provided
            "file_type": file_path.suffix.lower()[1:],
            "file_size": os.path.getsize(file_path),
            "created_time": os.path.getctime(file_path),
            "modified_time": os.path.getmtime(file_path),
            "schema_name": schema_name,
            "source_range_type": DocumentLoader.LOADER_CONFIGS[file_path.suffix.lower()]['source_range_type']
        }
        
        try:
            if file_path.suffix.lower() == '.pdf':
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    metadata.update({
                        "num_pages": len(pdf_reader.pages),
                        "pdf_info": pdf_reader.metadata or {},
                        "source_range": f"pages 1-{len(pdf_reader.pages)}"
                    })
            
            elif file_path.suffix.lower() == '.docx':
                doc = DocxDocument(file_path)
                metadata.update({
                    "num_paragraphs": len(doc.paragraphs),
                    "num_sections": len(doc.sections),
                    "word_properties": doc.core_properties.__dict__,
                    "source_range": f"paragraphs 1-{len(doc.paragraphs)}"
                })
            
            elif file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
                metadata.update({
                    "num_rows": len(df),
                    "num_columns": len(df.columns),
                    "column_names": list(df.columns),
                    "source_range": f"rows 1-{len(df)}"
                })
                
        except Exception as e:
            logger.warning(f"Error extracting detailed metadata for {file_path}: {str(e)}")
            metadata["metadata_error"] = str(e)
            
        return metadata

    @staticmethod
    def _extract_chunk_metadata(doc: LangchainDocument, base_metadata: Dict[str, Any], chunk_index: int) -> Dict[str, Any]:
        """
        Extract and enhance chunk-level metadata.
        
        Args:
            doc (LangchainDocument): Langchain document chunk
            base_metadata (Dict[str, Any]): Base document metadata
            chunk_index (int): Index of this chunk
            
        Returns:
            Dict[str, Any]: Enhanced chunk metadata
        """
        chunk_metadata = doc.metadata.copy()
        source_range_type = base_metadata["source_range_type"]
        
        # Add base metadata
        chunk_metadata.update({
            "document_id": base_metadata["document_id"],
            "schema_name": base_metadata.get("schema_name"),
            "chunk_index": chunk_index,
            "source_range_type": source_range_type
        })
        
        # Extract source range based on file type and metadata
        if source_range_type == 'page' and 'page' in chunk_metadata:
            chunk_metadata["source_range"] = f"page {chunk_metadata['page']}"
        elif source_range_type == 'row' and hasattr(doc, 'metadata') and 'row' in doc.metadata:
            chunk_metadata["source_range"] = f"row {doc.metadata['row']}"
        elif source_range_type == 'line':
            # For text files, use chunk index as line number
            chunk_metadata["source_range"] = f"line {chunk_index + 1}"
        elif source_range_type == 'paragraph':
            # For DOCX files, use chunk index as paragraph number
            chunk_metadata["source_range"] = f"paragraph {chunk_index + 1}"
        else:
            # Fallback to chunk index if no specific range can be determined
            chunk_metadata["source_range"] = f"chunk {chunk_index + 1}"
        
        return chunk_metadata

    @staticmethod
    def load_document(file_path: str, schema_name: Optional[str] = None, 
                     loader_config: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Load a document using Langchain document loaders and return its content and metadata.
        
        Args:
            file_path (str): Path to the document file
            schema_name (Optional[str]): Name of the schema this document belongs to
            loader_config (Optional[Dict[str, Any]]): Optional configuration overrides
            
        Returns:
            Dict[str, Any]: Dictionary containing:
                - content: str, the extracted text content
                - metadata: Dict containing filename, file_type, and format-specific metadata
                - chunks: List[Dict[str, Any]], individual chunks with their metadata
                - error: Optional[str], error message if loading failed
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                error_msg = f'File not found: {file_path}'
                return {"content": "", "metadata": {"error": error_msg}, "error": error_msg}
            
            file_type = file_path.suffix.lower()
            if file_type not in DocumentLoader.SUPPORTED_EXTENSIONS:
                error_msg = f'Unsupported file type: {file_type}'
                return {"content": "", "metadata": {"error": error_msg}, "error": error_msg}
            
            # Get default config and update with any overrides
            config = DocumentLoader.LOADER_CONFIGS[file_type].copy()
            if loader_config:
                config.update(loader_config)
            
            # Extract detailed metadata first
            metadata = DocumentLoader.extract_detailed_metadata(file_path, schema_name, config)
            
            # Try primary loader first
            try:
                result = DocumentLoader._load_with_langchain(file_path, config['primary_loader'], config, metadata)
            except Exception as primary_error:
                logger.warning(f"Primary loader failed for {file_path}: {str(primary_error)}")
                if 'fallback_loader' in config:
                    logger.info(f"Attempting fallback loader for {file_path}")
                    result = DocumentLoader._load_with_langchain(file_path, config['fallback_loader'], config, metadata)
                else:
                    raise primary_error
            
            # Update with detailed metadata
            result['metadata'].update(metadata)
            
            # Validate content
            validation_error = DocumentLoader.validate_content(result['content'], file_type)
            if validation_error:
                result['metadata']['validation_warning'] = validation_error
            
            return result
                
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            return {
                "content": "",
                "metadata": {"error": str(e)},
                "error": str(e)
            }
    
    @staticmethod
    def _load_with_langchain(file_path: Path, loader_class, config: Dict[str, Any], 
                            base_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Load a document using a Langchain document loader with configuration.
        
        Args:
            file_path (Path): Path to the document file
            loader_class: The Langchain loader class to use
            config (Dict[str, Any]): Loader configuration
            base_metadata (Dict[str, Any]): Base metadata to include
            
        Returns:
            Dict[str, Any]: Dictionary containing content, metadata, and chunks
        """
        try:
            # Initialize loader with format-specific configurations
            loader_kwargs = config.get('loader_kwargs', {}).copy()
            
            # Special handling for CSV files
            if file_path.suffix.lower() == '.csv':
                # CSVLoader has specific initialization requirements
                loader = loader_class(str(file_path))
            else:
                loader = loader_class(str(file_path), **loader_kwargs)
            
            # Load the document
            documents = loader.load()
            
            # Special handling for CSV files - each row becomes a chunk
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
                documents = []
                for idx, row in df.iterrows():
                    content = ", ".join(f"{col}: {val}" for col, val in row.items())
                    doc = LangchainDocument(
                        page_content=content,
                        metadata={"row": idx + 1}  # 1-based row numbers
                    )
                    documents.append(doc)
            
            # Process each chunk with enhanced metadata
            chunks = []
            for i, doc in enumerate(documents):
                chunk_metadata = DocumentLoader._extract_chunk_metadata(doc, base_metadata, i)
                chunks.append({
                    "content": doc.page_content,
                    "metadata": chunk_metadata
                })
            
            # Combine all document chunks into a single text
            content = "\n\n".join([chunk["content"] for chunk in chunks])
            
            # Update base metadata
            base_metadata['num_chunks'] = len(chunks)
            
            return {
                "content": content,
                "metadata": base_metadata,
                "chunks": chunks,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Error loading {file_path} with {loader_class.__name__}: {str(e)}")
            return {
                "content": "",
                "metadata": {"error": str(e)},
                "error": str(e)
            }
    
    @staticmethod
    def load_documents_from_directory(directory_path: str, schema_name: Optional[str] = None,
                                    recursive: bool = False, 
                                    loader_config: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory_path (str): Path to the directory
            schema_name (Optional[str]): Name of the schema these documents belong to
            recursive (bool): Whether to search subdirectories
            loader_config (Optional[Dict[str, Any]]): Optional configuration overrides
            
        Returns:
            List[Dict[str, Any]]: List of document dictionaries
        """
        directory = Path(directory_path)
        if not directory.exists() or not directory.is_dir():
            logger.error(f"Directory not found: {directory_path}")
            return []
        
        documents = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in DocumentLoader.SUPPORTED_EXTENSIONS:
                result = DocumentLoader.load_document(str(file_path), schema_name, loader_config)
                if result["error"] is None:
                    documents.append(result)
                else:
                    logger.warning(f"Skipping file {file_path} due to error: {result['error']}")
        
        return documents 