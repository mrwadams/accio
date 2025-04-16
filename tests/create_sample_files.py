from docx import Document
from docx.shared import Inches
import os
from pathlib import Path

# Create sample DOCX file
def create_sample_docx():
    doc = Document()
    doc.add_heading('Sample Document', 0)
    doc.add_paragraph('This is a sample document for testing the document loader.')
    doc.add_paragraph('It contains multiple paragraphs to test extraction.')
    doc.add_paragraph('This is the third paragraph.')
    
    # Save the document
    doc.save('tests/data/sample.docx')
    print("Created sample DOCX file")

if __name__ == "__main__":
    # Ensure the data directory exists
    os.makedirs('tests/data', exist_ok=True)
    
    # Create sample files
    create_sample_docx()
    
    print("All sample files created successfully!")
    print("Note: sample.pdf and sample.csv should be manually added to tests/data/") 